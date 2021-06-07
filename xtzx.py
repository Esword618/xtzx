# 学堂账号密码登录
import os
import random
import re
import string
import sys
import requests
from tqdm import tqdm
import json
from requests.utils import dict_from_cookiejar, cookiejar_from_dict
from websocket import create_connection
import time
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

"""
错误处理函数
"""


def error(error_path):
    print(f'程序在{error_path}有bug\n'
          f'程序自动结束')
    sys.exit()


"""
登录函数
"""


class Login():

    def __init__(self):
        self.session = requests.session()

    def code_img(self, img_url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36',
        }
        response = requests.get(url=img_url, headers=headers)
        content = response.content
        with open('core/QR.png', 'wb') as f:
            f.write(content)
        base_path = os.getcwd()
        os.startfile(f'{base_path}\core/QR.png')
        return

    def get_data(self):
        print('data')
        headers = {
            'Connection': 'Upgrade',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36',
            'Host': 'www.xuetangx.com',
            'Origin': 'https://www.xuetangx.com',
            'Sec-WebSocket-Key': 'jv7GrgzlhDyfGs3H7TlWGw==',
            'Sec-WebSocket-Version': '13',
            'Upgrade': 'websocket',
        }
        url = "wss://www.xuetangx.com/wsapp/"
        wss = create_connection(url=url, header=headers)
        '''
        payload必须json化,否则报错
        '''
        payload = {'op': 'requestlogin', 'purpose': 'login', 'role': 'web', 'version': '1.4'}
        payload = json.dumps(payload)
        wss.send(payload)
        base_code_url = json.loads(wss.recv())['ticket']
        self.code_img(base_code_url)
        while True:
            print('请扫码!')
            wss.send(payload)
            response_text = wss.recv()
            if 'token' in response_text:
                print('扫码成功!')
                data = json.loads(response_text)
                os.remove('core/QR.png')
                return data
            time.sleep(0.1)

    def save_cookie(self):
        cookie_dic = dict_from_cookiejar(self.session.cookies)
        with open('cookie/XTcookies.json', 'w') as f:
            f.write(json.dumps(cookie_dic))
        print('学堂在线cookie保存成功！')

    def reade_cookie(self):
        b = os.path.exists("cookie/XTcookies.json")
        if b == False:
            print("没有cookie文件")
            return False
        else:
            with open(r'cookie/XTcookies.json', 'r') as f:
                cookie_dic = json.loads(f.read())
            cookie_jar = cookiejar_from_dict(cookie_dic)
            self.session.cookies = cookie_jar
            print("cookie已加载")
            return True

    def check_cookie(self):
        headers = {
            'referer': 'https://www.xuetangx.com/',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36',
            'x-client': 'web',
            'x-requested-with': 'XMLHttpRequest',
            'xtbz': 'xt',
        }
        url = 'https://www.xuetangx.com/api/v1/u/login/check_is_l/'
        response = self.session.get(url=url, headers=headers)
        is_login = response.json()['data']['is_login']
        if is_login == False:
            return False
        else:
            return True

    def login(self):
        data = self.get_data()
        payload = {
            's_s': data['token'],
        }
        url = 'https://www.xuetangx.com/api/v1/u/login/wx/'
        response = self.session.post(url=url, json=payload)
        status_code = response.status_code
        return status_code

    def main(self):
        b = self.reade_cookie()
        if b == False:
            b = self.check_cookie()
            if b == False:
                status_code = self.login()
                if status_code == 200:
                    print('登录成功！')
                    self.save_cookie()
        else:
            b = self.check_cookie()
            if b == False:
                status_code = self.login()
                if status_code == 200:
                    print('登录成功！')
                    self.save_cookie()
        return self.session


class Spider():
    """
    初始化参数
    """

    def __init__(self):
        self.session = Login().main()
        self.session.headers = {
            'referer': 'https://www.xuetangx.com/',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36',
            'x-client': 'web',
            'x-csrftoken': self.get_csrftoken(),
            'xtbz': 'xt',
        }

    """
    获取x-csrftoken,x-csrftoken在cookies里面
    """

    def get_csrftoken(self):
        cookie_dict = dict_from_cookiejar(self.session.cookies)
        csrftoken = cookie_dict['csrftoken']
        return csrftoken

    """
    获取课程页数，返回类型为整数
    """

    def get_max_page(self):
        url = 'https://www.xuetangx.com/api/v1/lms/suggest_keyword/?keyword='
        response = self.session.get(url=url)
        max_page = response.json()['data']['max_page']
        return max_page

    """
    get_courses为获取课程信息
    max_page为自己所拥有课程的最大页数
    """

    def get_courses(self, max_page):
        url = 'https://www.xuetangx.com/api/v1/lms/user/user-courses/'
        """
        dict_courses为一个字典，里面含有course的info
        """
        self.dict_courses = {}
        dict_courses = {}
        num = 0
        for page in range(1, max_page + 1):
            params = {
                'status': '1',
                'page': f'{page}',
            }
            response = self.session.get(url=url, params=params)
            if response.status_code == 200:
                data = response.json()['data']
                product_list = data['product_list']
                for product in product_list:
                    """
                    name为课程名字
                    classroom_id为课程id
                    sign为课程sign
                    """
                    name = product['name']
                    print(num,name)
                    self.dict_courses[str(num)] = name
                    num += 1
                    classroom_id = product['classroom_id']
                    sign = product['sign']
                    """
                    dict_courses = {{'课程1':{'sign':sign,'classroom_id':classroom_id}},{'课程2':{'sign':sign,'classroom_id':classroom_id}}........}
                    """
                    dict_courses[name] = {}
                    dict_courses[name]['sign'] = sign
                    dict_courses[name]['classroom_id'] = classroom_id
            else:
                error_path = '学堂在线get_courses'
                error(error_path)
        return dict_courses

    '''试卷'''
    '''------------------------------------------------------------------------------------------'''
    """
    get_paper_params函数为
    获取作业name及其id
    classroom_id为params中的cid
    """

    def get_paper_params(self, dict_course):
        dict_section = {}
        cid = dict_course['classroom_id']
        sign = dict_course['sign']
        self.classroom_id = cid
        self.sign = sign

        url = 'https://www.xuetangx.com/api/v1/lms/learn/course/chapter'
        params = {
            'cid': cid,
            'sign': sign,
            'etag_id': '11',
        }
        response = self.session.get(url=url, params=params)
        data = response.json()['data']

        """
        course_chapter为一个列表,如一个单元语文标题
        section_leaf_list为一个单元里面的课文题目
        """

        course_chapter = data['course_chapter']
        for course in course_chapter:
            section_leaf_list = course['section_leaf_list']
            for section in section_leaf_list:

                """
                这里会遇到2种情况:
                情况1:
                {
                  "order": 0,
                  "leaf_list": [
                    {
                      "name": "（a）计算--作业",
                      "is_locked": false,
                      "start_time": 1609689600000,
                      "chapter_id": 958633,
                      "section_id": 2430690,
                      "leaf_type": 6,
                      "id": 9215536,
                      "is_show": true,
                      "end_time": 0,
                      "score_deadline": 1627228799000,
                      "is_score": true,
                      "is_assessed": false,
                      "order": 7,
                      "leafinfo_id": 9218119
                    }
                  ],
                  "chapter_id": 958633,
                  "id": 2430690,
                  "name": "（a）计算"
                }
                情况2:
                {
                  "name": "（b）计算模型--作业",
                  "is_locked": false,
                  "start_time": 1609689600000,
                  "chapter_id": 958633,
                  "section_id": null,
                  "leaf_type": 6,
                  "id": 9215552,
                  "is_show": true,
                  "end_time": 0,
                  "score_deadline": 1627228799000,
                  "is_score": true,
                  "is_assessed": false,
                  "order": 2,
                  "leafinfo_id": 9218135
                }
                """

                if 'leaf_list' in section:
                    leaf_list = section['leaf_list']
                    for leaf in leaf_list:
                        name = leaf['name']
                        if '本章测验--作业' in name:
                            name = name + string.digits[self.name_num]
                            self.name_num += 1
                        id = leaf['id']
                        dict_section[name] = id
                else:
                    name = section['name']
                    if '本章测验--作业' in name:
                        name = name + string.digits[self.name_num]
                        self.name_num += 1
                    id = section['id']
                    dict_section[name] = id
        return dict_section

    """
    dict_section处理函数
    """

    def handle_dict_section(self, dict_section):
        for name, leaf_id in dict_section.items():
            print(f'正在进行{name}')
            self.name = name
            self.leaf_id = leaf_id
            self.exercise_id = self.get_exercise_id()
            self.get_paper_info()
            self.word()

    def handle_re_html(self, html):
        list = []
        html = re.sub(r'&nbsp;|\n', '', html)
        html_list = re.findall('src="(.*?)"|>(.*?)<', html, re.S)
        for i, j in html_list:
            if i != '':
                list.append(i)
            if j != '':
                list.append(j)
        return list

    """
    explain:答案解析
    content:题目信息
    Options:选项
    """

    def handle_topic(self, TypeText, Body, explain, Options):
        # topic_dict = {}
        """题头信息提取"""
        option_dict = {}
        Body = self.handle_re_html(Body)
        """解析信息提取"""
        if explain != None:
            explain = self.handle_re_html(explain)

        # 选择题and判断题
        if '选' in TypeText:
            for Option in Options:
                key = Option['key']
                value = Option['value']
                """选项信息提取"""
                value = self.handle_re_html(value)
                option_dict[key] = value
        elif TypeText == '判断题':
            for Option in Options:
                key = Option['key']
                value = ''
                """选项信息提取"""
                option_dict[key] = value
        elif TypeText == '主观题':
            pass

        elif TypeText == '填空题':
            pass

        return Body, explain, option_dict

    """
    handle_TypeText:题目类型判断处理函数
    """

    def handle_TypeText(self, answer_count, problems_count, problem):
        """
        user:用户信息
        content:题目主体
        TypeText:试题类型
        Body:题头
        explain:解释
        """
        explain = None
        Options = None
        user = problem['user']
        if user['is_show_explain'] == True:
            explain = user['explain']
            if explain == [] or explain == '' or explain == ' ':
                explain = None
        content = problem['content']
        TypeText = content['TypeText']
        Body = content['Body']
        answer = ''
        """
        下面是一个试题类型判断
        """
        # 选择题
        if '选' in TypeText or TypeText == '判断题':
            """
            Options:选项列表
            """
            Options = content['Options']

            if answer_count == problems_count:
                answers = problem['user']['answer']
                if type(answers) == list:
                    for i in answers:
                        answer += i
                else:
                    answer = answers
                # 可以直接获取试题
            else:
                is_show_answer = user['is_show_answer']
                if is_show_answer == False:
                    answer = self.upload_answer(TypeText=TypeText)
                else:
                    answer_list = user['answer']
                    for i in answer_list:
                        answer += i

        # elif TypeText == '主观题':
        #     if answer_count == problems_count:
        #         if 'answer' in user.keys():
        #             answer = user['answer']
        #             answer = self.handle_re_html(answer)
        #         else:
        #             answer = '暂时无答案'
        #         # 可以直接获取试题
        #     else:
        #         is_show_answer = user['is_show_answer']
        #         if is_show_answer == False:
        #             answer = self.upload_answer(TypeText=TypeText)
        #         else:
        #             if 'answer' in user.keys():
        #                 answer = user['answer']
        #                 answer = self.handle_re_html(answer)
        #             else:
        #                 answer = '暂时无答案'
        elif TypeText == '填空题':
            count = user['count']
            if answer_count == problems_count:
                answers = user['answers']
                for i in answers.values():
                    for j in i:
                        answer += f'{j},'
            else:
                is_show_answer = user['is_show_answer']
                if is_show_answer == False:
                    answer = self.upload_answer(TypeText=TypeText, count=count)
                else:
                    if 'answers' in user.keys():
                        answers = user['answers']
                        for i in answers.values():
                            for j in i:
                                answer += f'{j},'
                    else:
                        answer = '暂时无答案'

        Body, explain, option_dict = self.handle_topic(TypeText=TypeText, Body=Body, explain=explain, Options=Options)

        return Body, explain, option_dict, answer

    """
    get_exercise_id:获取exercise_id
    """

    def get_exercise_id(self):
        url = f'https://www.xuetangx.com/api/v1/lms/learn/leaf_info/{self.classroom_id}/{self.leaf_id}/?sign={self.sign}'
        response = self.session.get(url=url)
        exercise_id = response.json()['data']['content_info']['leaf_type_id']
        return exercise_id

    """
    get_paper_info函数为获取作业的题目并正好为试卷
    """

    def get_paper_info(self):
        self.word_dict = {}
        url = f'https://www.xuetangx.com/api/v1/lms/exercise/get_exercise_list/{self.exercise_id}/'
        response = self.session.get(url=url)
        data = response.json()['data']
        answer_count = data['answer_count']
        problems = data['problems']
        problems_count = len(problems)

        """
        problems为一个试题列表
        problem为一道试题(填空/选择)
        """
        for problem in problems:
            self.problem_id = problem['problem_id']
            Body, explain, option_dict, answer = self.handle_TypeText(answer_count=answer_count,
                                                                      problems_count=problems_count, problem=problem)
            Body = tuple(Body)
            self.word_dict[Body] = {}
            self.word_dict[Body]['options'] = option_dict
            self.word_dict[Body]['answer'] = answer
            self.word_dict[Body]['explain'] = explain

    def save_img(self, url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
        }
        response = requests.get(url=url, headers=headers)
        content = response.content
        with open('core\paper.png', 'wb') as f:
            f.write(content)
        return

    def word(self):
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        num = 1
        for Body, value in tqdm(self.word_dict.items(), desc=self.name):
            answer = value['answer']
            option_dict = value['options']
            explain = value['explain']
            '''开启一个新段落'''
            paragraph = document.add_paragraph(f'{num},')
            num += 1
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            run = document.paragraphs[-1].add_run()
            for i in Body:
                if 'http' in i:
                    try:
                        url = i.replace('  ', '')
                        self.save_img(url)
                        run.add_picture('core\paper.png')
                    except:
                        run.add_text(u'{}'.format("此处可能有照片，但无法正常显示"))
                else:
                    run.add_text(u'{}'.format(i))

            if option_dict != {}:
                for key, value in option_dict.items():
                    '''key代表选项ABCD,value代表选项内容，是一个列表'''
                    '''开启一个新段落'''
                    paragraph = document.add_paragraph(f'{key},')
                    # 设置行距
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.5  # 1.5倍行距
                    run = document.paragraphs[-1].add_run()
                    for i in value:
                        if 'http' in i:
                            url = i.replace('  ', '')
                            self.save_img(url)
                            run.add_picture('core\paper.png')
                        else:
                            run.add_text(u'{}'.format(i))
            '''答案'''
            '''开启一个新段落'''
            paragraph = document.add_paragraph(f'答案:{answer}')
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距

            try:
                if explain != [] or explain != '' or explain != None:
                    document.add_paragraph('解析:')
                    paragraph = document.add_paragraph('')
                    # 设置行距
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.5  # 1.5倍行距
                    run = document.paragraphs[-1].add_run()
                    for i in explain:
                        if 'http' in i:
                            url = i.replace('  ', '')
                            self.save_img(url)
                            run.add_picture('core\paper.png')
                        else:
                            run.add_text(u'{}'.format(i))
            except:
                pass
        document.save('{}/{}.docx'.format(self.path, self.name))

    """
    上传答案
    """

    def upload_answer(self, TypeText, count=None):
        time.sleep(1.5)
        url = 'https://www.xuetangx.com/api/v1/lms/exercise/problem_apply/'
        answer = ''
        if '选' in TypeText:
            payload = {
                "leaf_id": self.leaf_id,
                "classroom_id": self.classroom_id,
                "exercise_id": self.exercise_id,
                "problem_id": self.problem_id,
                "sign": self.sign,
                "answers": {},
                "answer": ["A"]
            }
            response = self.session.post(url=url, json=payload)
            data = response.json()['data']
            answer_list = data['answer']
            for i in answer_list:
                answer += i
            time.sleep(0.75)

        elif TypeText == '判断题':
            payload = {
                "leaf_id": self.leaf_id,
                "classroom_id": self.classroom_id,
                "exercise_id": self.exercise_id,
                "problem_id": self.problem_id,
                "sign": self.sign,
                "answers": {},
                "answer": random.choice([["true"], ["false"]])
            }
            response = self.session.post(url=url, json=payload)
            data = response.json()['data']
            answer_list = data['answer']
            for i in answer_list:
                answer += i
            time.sleep(1.5)

        # elif TypeText == '主观题':
        #     payload = {
        #         "leaf_id": self.leaf_id,
        #         "classroom_id": self.classroom_id,
        #         "exercise_id": self.exercise_id,
        #         "problem_id": self.problem_id,
        #         "sign": self.sign,
        #         "answers": {},
        #         "answer": {
        #             "time": "0",
        #             "content": "<div class=\"custom_ueditor_cn_body\"><p>我真的帅</p></div>",
        #             "oSubject": {"attachments": {"filelist": []}}}
        #     }
        #     response = self.session.post(url=url, json=payload)
        #     data = response.json()['data']
        #     if 'answer' in data.keys():
        #         answer = data['answer']
        #         answer = self.handle_re_html(answer)
        #     else:
        #         answer = '暂时无答案'

        elif TypeText == '填空题':
            answers = {f"{i}": "1" for i in range(1, count + 1)}
            payload = {
                "leaf_id": self.leaf_id,
                "classroom_id": self.classroom_id,
                "exercise_id": self.exercise_id,
                "problem_id": self.problem_id,
                "sign": self.sign,
                "answers": answers,
                "answer": ""}
            response = self.session.post(url=url, json=payload)
            data = response.json()['data']
            if 'answers' in data.keys():
                answers = data['answers']
                for i in answers.values():
                    try:
                        for j in i:
                            answer += j
                    except:
                        answer = i[0]
            else:
                answer = '暂时无答案'

        return answer

    "------------------------------------------------------------------------------"

    def get_movie_response_dict_name_url(self, dict_course):
        dict_name_url = {}
        classroom_id = dict_course['classroom_id']
        sign = dict_course['sign']
        url = 'https://www.xuetangx.com/api/v1/lms/learn/course/chapter'
        params = {
            'cid': classroom_id,
            'sign': sign,
            'etag_id': '6',
        }
        response = self.session.get(url=url, params=params)
        course_chapter = response.json()['data']['course_chapter']
        xmind = ''
        print('-------------------')
        for course in course_chapter:
            name = course['name']
            xmind += f'|---》》》{name}\n'
            section_leaf_list = course['section_leaf_list']
            for section_leaf in section_leaf_list:
                leaf_list = section_leaf['leaf_list']
                xmind += f'|------》》{name}\n'
                name = section_leaf['name']
                for leaf in leaf_list:
                    name = leaf['name']
                    xmind += f'|----------》{name}\n'
                    id = leaf['id']
                    url = f'https://www.xuetangx.com/api/v1/lms/learn/leaf_info/{classroom_id}/{id}/?sign={sign}'
                    dict_name_url[name] = url

        print(xmind)
        return dict_name_url

    def get_movie_ccid(self, url):
        ccid = self.session.get(url=url).json()['data']['content_info']['media']['ccid']
        return ccid

    def get_movie_url(self, ccid, video_quality):
        video_quality = {0: 'quality10', 1: 'quality20'}[video_quality]
        url = f'https://www.xuetangx.com/api/v1/lms/service/playurl/{ccid}/?appid=10000'
        response = self.session.get(url=url)
        movie_url = response.json()['data']['sources'][video_quality][0]
        return movie_url

    def download_movie(self, movie_url, path, name):
        headers = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36',
        }
        response = requests.get(url=movie_url, headers=headers)
        temp_size = 0  # 已经下载文件大小
        chunk_size = 1024  # 每次下载数据大小
        start = time.time()
        total_size = int(response.headers.get("Content-Length"))
        with open(path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=chunk_size):
                if chunk:
                    temp_size += len(chunk)
                    f.write(chunk)
                    f.flush()
                    #############花哨的下载进度部分###############
                    done = int(50 * temp_size / total_size)
                    # 调用标准输出刷新命令行，看到\r 回车符了吧
                    # 相当于把每一行重新刷新一遍
                    sys.stdout.write(
                        "\r[%s%s] %d%%" % ('█' * done, ' ' * (50 - done), 100 * temp_size / total_size))
                    sys.stdout.flush()
        print()  # 避免上面\r 回车符，执行完后需要换行了，不然都在一行显示
        end = time.time()  # 结束时间
        print(f'{name}下载完成!用时%.2f 秒' % (end - start))

    '''--------------------------------------------------------------------'''
    '''图文资料'''

    def get_data_response_dict_name_url(self, dict_course):
        dict_name_url = {}
        classroom_id = dict_course['classroom_id']
        sign = dict_course['sign']
        url = 'https://www.xuetangx.com/api/v1/lms/learn/course/chapter'
        params = {
            'cid': classroom_id,
            'sign': sign,
            'etag_id': '7',
        }
        response = self.session.get(url=url, params=params)
        course_chapter = response.json()['data']['course_chapter']
        xmind = ''
        print('-------------------')
        for course in course_chapter:
            name = course['name']
            xmind += f'|---》》》{name}\n'
            section_leaf_list = course['section_leaf_list']
            for section_leaf in section_leaf_list:
                leaf_list = section_leaf['leaf_list']
                xmind += f'|------》》{name}\n'
                name = section_leaf['name']
                for leaf in leaf_list:
                    name = leaf['name']
                    xmind += f'|----------》{name}\n'
                    id = leaf['id']
                    url = f'https://www.xuetangx.com/api/v1/lms/learn/leaf_info/{classroom_id}/{id}/?sign={sign}'
                    dict_name_url[name] = url

        print(xmind)
        return dict_name_url

    def get_data_url_name(self, url):
        response = self.session.get(url=url)
        download = response.json()['data']['content_info']['download']
        if download != []:
            download = download[0]
            file_name = download['file_name']
            file_name = re.sub('\s+|:|：', '', file_name)
            if '.pdf' not in file_name:
                file_name = file_name + '.pdf'
            file_url = download['file_url']
            return file_name, file_url
        else:
            return False

    def download_data(self, file_url, path, name):
        headers = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36',
        }
        response = requests.get(url=file_url, headers=headers)
        temp_size = 0  # 已经下载文件大小
        chunk_size = 1024  # 每次下载数据大小
        start = time.time()
        total_size = int(response.headers.get("Content-Length"))
        with open(path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=chunk_size):
                if chunk:
                    temp_size += len(chunk)
                    f.write(chunk)
                    f.flush()
                    #############花哨的下载进度部分###############
                    done = int(50 * temp_size / total_size)
                    # 调用标准输出刷新命令行，看到\r 回车符了吧
                    # 相当于把每一行重新刷新一遍
                    sys.stdout.write(
                        "\r[%s%s] %d%%" % ('█' * done, ' ' * (50 - done), 100 * temp_size / total_size))
                    sys.stdout.flush()
        print()  # 避免上面\r 回车符，执行完后需要换行了，不然都在一行显示
        end = time.time()  # 结束时间
        print(f'{name}_____下载完成!用时%.2f 秒' % (end - start))

    '''--------------------------------------------------------------------'''

    def main(self):
        """
        主函数
        """
        max_page = self.get_max_page()
        dict_courses = self.get_courses(max_page=max_page)
        while True:
            self.name_num = 0
            num = input('请输入课程序号:')
            course_name = self.dict_courses[num]
            dict_course = dict_courses[course_name]
            self.path = f'data/{course_name}'
            b = os.path.exists(self.path)
            if b == False:
                os.makedirs(self.path)
            '''试卷'''
            def papaer():
                dict_section = self.get_paper_params(dict_course=dict_course)
                self.handle_dict_section(dict_section=dict_section)
            '''视频'''
            def movie():
                dict_name_url = self.get_movie_response_dict_name_url(dict_course=dict_course)
                video_quality = int(input('0代表标清\n'
                                          '1代表高清\n'
                                          '请输入视频清晰度:'))
                print('开始下载......')
                for name,url in dict_name_url.items():
                    name = re.sub('\s+|:|：', '', name)
                    path = fr'{self.path}/{name}.mp4'
                    if os.path.exists(path) == False:
                        ccid = self.get_movie_ccid(url=url)
                        movie_url = self.get_movie_url(ccid=ccid,video_quality=video_quality)
                        self.download_movie(movie_url=movie_url,path=path,name=name)
                    else:
                        print(f'{name}.mp4___已经存在!')
            '''图文资料'''
            def data():
                dict_name_url = self.get_data_response_dict_name_url(dict_course=dict_course)
                for name, url in dict_name_url.items():
                    b = self.get_data_url_name(url=url)
                    if b != False:
                        file_name, file_url = b
                        path = f'{self.path}\{file_name}'
                        if os.path.exists(path) == False:
                            self.download_data(file_url=file_url, path=path, name=name)
                        else:
                            print(f'{name}_____已存在!')
                    else:
                        print(f'{name}______为无可下载文件!')
            while True:
                print('0代表试卷\n'
                      '1代表视频\n'
                      '2代表图文资料\n'
                      '3代表全部\n'
                      '其它代表上一级\n')
                num = input('请输入数字:')
                if num == '0':
                    papaer()
                elif num == '1':
                    movie()
                elif num == '2':
                    data()
                elif num == '3':
                    papaer()
                    movie()
                    data()
                else:
                    break

#
# d = Spider()
# d.main()



def main():
    if os.path.exists('core') == False:
        os.makedirs('core')
    if os.path.exists('data') == False:
        os.makedirs('data')
    if os.path.exists('cookie') == False:
        os.makedirs('cookie')
    d = Spider()
    d.main()
    try:
        os.remove("core/paper.png")
    except:
        pass

if __name__ == '__main__':
    main()


"""
Spider参数获取流程
get_max_page:首先获取自己课程所占页数
get_courses:获取课程 name 以及课程 classroom_id与sign;
get_courses返回一个字典，{name:{'classroom_id:'classroom_id,'sign':sign}......}
然后选择课程，通过课程 name 获取上面字典与之对应的 classroom_id与sign
然后把获取到的value字典传入get_paper_params
get_paper_params:获取试卷参数
"""